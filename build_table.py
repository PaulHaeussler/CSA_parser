import docx
from colorama import Fore, Back, Style, init
from datetime import datetime

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


from docx.shared import Pt, Cm, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.table import _Cell

from csa_parser import CSA

init(convert=True)

def set_border(cell):
    set_cell_border(
        cell,
        top={"sz": 6, "val": "single", "color": "#BFBFBF"},
        bottom={"sz": 6, "val": "single", "color": "#BFBFBF"},
        start={"sz": 6, "val": "single", "color": "#BFBFBF"},
        end={"sz": 6, "val": "single", "color": "#BFBFBF"},
    )

def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))





# Creating a new file to generate the table into
csa = CSA()

print(f"{Fore.LIGHTCYAN_EX}Creating word table...{Style.RESET_ALL}")


doc = docx.Document()
# Creating filename with full length date string to avoid collisions (checking not necessary)
filename = f"CSA Table {csa.sections[0].subsections[0].questions[0].answer} {datetime.now().strftime('%d%m%Y %H-%M-%S')}.docx"
# Save empty workbook initially to establish name is not used and getting a file handle
doc.save(filename)


QUESTION = "Frage"
ANSWER = "Antwort"
COMMENT = "Kommentar"

if csa.language == "EN":
    QUESTION = "Question"
    ANSWER = "Answer"
    COMMENT = "Comment"

doc.add_paragraph().add_run(f"Questions answered: {csa.questions_answered}")
doc.add_paragraph().add_run(f"Anzahl Tables: {csa.ss_count}")
sc = 1  # section counter
for s in csa.sections:
    print(f"  Creating section {sc} {s.name}")
    # Set title
    pt = doc.add_paragraph().add_run(f"{sc}. {s.name}")
    pt.font.size = Pt(12)
    pt.font.bold = True
    pt.font.name = "Arial"

    ssc = 1  # subsection counter
    for ss in s.subsections:
        t = doc.add_table(rows=2,  cols=3)
        # Subsection Header
        t.cell(0, 0).merge(t.cell(0, 2)).text = f"{sc}.{ssc} {ss.name}"
        t.cell(0, 0).height = Cm(0.42)
        t.cell(0, 0).paragraphs[0].runs[0].font.bold = True
        t.cell(0, 0).paragraphs[0].runs[0].font.size = Pt(11)
        t.cell(0, 0).paragraphs[0].runs[0].font.name = "Arial"
        set_border(t.cell(0, 0))

        t.cell(1, 0).width = Cm(1.3)
        t.cell(1, 0).height = Cm(0.42)
        set_border(t.cell(1, 0))
        # Question Header
        t.cell(1, 1).width = Cm(8.0)
        t.cell(1, 1).height = Cm(0.42)
        t.cell(1, 1).text = QUESTION
        t.cell(1, 1).paragraphs[0].runs[0].font.bold = True
        t.cell(1, 1).paragraphs[0].runs[0].font.size = Pt(11)
        t.cell(1, 1).paragraphs[0].runs[0].font.name = "Arial"
        t.cell(1, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 157, 224)
        set_border(t.cell(1, 1))
        # Answer header
        t.cell(1, 2).width = Cm(7.2)
        t.cell(1, 2).height = Cm(0.42)
        t.cell(1, 2).text = ANSWER
        t.cell(1, 2).paragraphs[0].runs[0].font.bold = True
        t.cell(1, 2).paragraphs[0].runs[0].font.size = Pt(11)
        t.cell(1, 2).paragraphs[0].runs[0].font.name = "Arial"
        t.cell(1, 2).paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 157, 224)
        set_border(t.cell(1, 2))

        for q in ss.questions:
            r = t.add_row().cells
            r[0].text = q.index
            r[0].width = Cm(1.3)
            r[0].height = Cm(0.42)
            r[0].paragraphs[0].runs[0].font.bold = True
            r[0].paragraphs[0].runs[0].font.size = Pt(11)
            r[0].paragraphs[0].runs[0].font.name = "Arial"
            set_border(r[0])
            r[1].text = q.question
            r[1].width = Cm(8.0)
            r[1].height = Cm(0.42)
            r[1].paragraphs[0].runs[0].font.size = Pt(11)
            r[1].paragraphs[0].runs[0].font.name = "Arial"
            set_border(r[1])
            if q.type != "M" and q.type != "E":
                r[2].text = q.answer
                r[2].width = Cm(7.2)
                r[2].height = Cm(0.42)
                r[2].paragraphs[0].runs[0].font.size = Pt(11)
                r[2].paragraphs[0].runs[0].font.name = "Arial"
                set_border(r[2])
            else:
                r[1].merge(r[2])
            if q.type == "M" or q.type == "E":
                for i, key in enumerate(q.options):
                    rr = t.add_row().cells
                    rr[0].width = Cm(1.3)
                    rr[0].height = Cm(0.42)
                    set_border(rr[0])
                    rr[1].text = key
                    rr[1].width = Cm(14.2)
                    rr[1].height = Cm(0.42)
                    rr[1].paragraphs[0].runs[0].font.size = Pt(11)
                    rr[1].paragraphs[0].runs[0].font.name = "Arial"
                    set_border(rr[1])
                    if q.type == "M":
                        if q.options[key]:
                            rr[2].text = "X"
                        else:
                            rr[2].text = ""
                    else:
                        rr[2].text = q.options[key]
                    rr[2].width = Cm(1.0)
                    rr[2].height = Cm(0.42)
                    rr[2].paragraphs[0].runs[0].font.size = Pt(11)
                    rr[2].paragraphs[0].runs[0].font.name = "Arial"
                    set_border(rr[2])


            # Comment
            if hasattr(q, "comment"):
                cr = t.add_row().cells
                cr[0].width = Cm(1.3)
                cr[0].height = Cm(0.42)
                set_border(cr[0])
                cr[1].width = Cm(8.0)
                cr[1].height = Cm(0.42)
                cr[1].text = COMMENT
                cr[1].paragraphs[0].runs[0].font.size = Pt(11)
                cr[1].paragraphs[0].runs[0].font.name = "Arial"
                set_border(cr[1])
                cr[2].width = Cm(7.2)
                cr[2].height = Cm(0.42)
                cr[2].text = q.comment
                cr[2].paragraphs[0].runs[0].font.size = Pt(11)
                cr[2].paragraphs[0].runs[0].font.name = "Arial"
                set_border(cr[2])

        doc.add_paragraph().add_run("")
        doc.add_paragraph().add_run("")
        ssc += 1
    sc += 1


doc.save(filename)
print(f"{Fore.LIGHTGREEN_EX}Successfully finished creating the table!{Style.RESET_ALL}")


